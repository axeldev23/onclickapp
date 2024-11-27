from rest_framework import viewsets
from rest_framework.response import Response
from rest_framework.decorators import api_view
from .models import Prestamo, Cliente, Pago
from .serializer import PrestamoSerializer, ClienteSerializer, UserSerializer, PagoSerializer
from django.contrib.auth.models import User
from rest_framework.authtoken.models import Token
from rest_framework import status
from django.shortcuts import get_object_or_404
from rest_framework.decorators import permission_classes, authentication_classes, parser_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.authentication import TokenAuthentication
from django.http import Http404
from django.http import FileResponse
from rest_framework.parsers import MultiPartParser, FormParser
import io
from django.conf import settings
import os
from docx import Document
from django.http import HttpResponse
from .serializer import DocumentSerializerAmortizacion, DocumentSerializer
from rest_framework.views import APIView
from django.utils.decorators import method_decorator
from django.views.decorators.csrf import csrf_exempt
from python_docx_replace import docx_replace
from datetime import datetime, timedelta
import requests
from django.utils import timezone
from twilio.rest import Client
from decouple import config




# FORMATO TABLA DE AMORTIZACIÓN
@method_decorator(csrf_exempt, name='dispatch')
class AmortizacionAPIView(APIView):
    def post(self, request):
        serializer = DocumentSerializerAmortizacion(data=request.data)
        if serializer.is_valid():
            try:
                doc_path = os.path.join(settings.BASE_DIR, 'static', 'formatos', 'TablaAmortizacionFormato.docx')
                doc = Document(doc_path)

                replacements = {
                    'clientes.nombre_completo': serializer.validated_data['nombre_completo'],
                    'prestamos.equipo_a_adquirir': serializer.validated_data['equipo_a_adquirir'],
                    'prestamos.equipo_precio': str(serializer.validated_data['equipo_precio']),
                    'prestamos.pago_inicial': str(serializer.validated_data['pago_inicial']),
                    'prestamos.monto_credito': str(serializer.validated_data['monto_credito']),
                    'prestamos.plazo_credito': str(serializer.validated_data['plazo_credito']),
                    'variable_monto_parcialidad': str(serializer.validated_data['monto_parcialidad']),
                    'prestamos.total_a_pagar': str(serializer.validated_data['total_a_pagar']),
                    'prestamos.fecha_inicio': str(serializer.validated_data['fecha_inicio']),
                    'clientes.domicilio_actual': str(serializer.validated_data['domicilio_actual']),
                    'clientes.numero_telefono': serializer.validated_data['numero_telefono'],
                    'prestamos.prestamo_id': serializer.validated_data['prestamo_id'],
                    'prestamos.imei': serializer.validated_data['imei'],
                }

                # Reemplazar los campos de texto con sus valores
                docx_replace(doc, **replacements)

                # Encontrar la tabla específica (en este caso, la primera tabla)
                table = doc.tables[0]

                num_pagos = serializer.validated_data['plazo_credito']
                monto_pago = float(serializer.validated_data['monto_parcialidad'])
                fecha_inicial = serializer.validated_data['fecha_primer_pago']
                fecha_inicial_dt = datetime.strptime(fecha_inicial.strftime('%Y-%m-%d'), '%Y-%m-%d')

                for i in range(num_pagos):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(i + 1)
                    row_cells[1].text = (fecha_inicial_dt + timedelta(weeks=i)).strftime('%d-%m-%Y')
                    row_cells[2].text = f'${monto_pago:.2f}'
                    row_cells[3].text = ''  # Estado inicial de pago

                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="tabla_amortizacion_{serializer.validated_data["nombre_completo"]}.docx"'

                return response
            except Exception as e:
                return Response({'error': 'Error al procesar el documento.', 'details': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)




# FORMATO PAGARÉ
@method_decorator(csrf_exempt, name='dispatch')  # Deshabilitar la verificación CSRF
class DocumentAPIView(APIView):

    def post(self, request):
        serializer = DocumentSerializer(data=request.data)
        if serializer.is_valid():
            try:
                doc_path = os.path.join(settings.STATIC_ROOT, 'formatos', 'PAGARÉ Formato.docx')
                doc = Document(doc_path)
                
                replacements = {
                    'prestamos.fecha_inicio': serializer.validated_data['fecha_inicio'],
                    'cliente.nombre_completo': serializer.validated_data['nombre_completo'],
                    'cliente.clave_elector': serializer.validated_data['clave_elector'],
                    'cliente.domicilio_actual': serializer.validated_data['domicilio_actual'],
                    'variable_total_a_pagar': serializer.validated_data['variable_total_a_pagar'],
                    'prestamos.equipo_a_adquirir': serializer.validated_data['equipo_a_adquirir'],
                    'prestamos.interes': serializer.validated_data['interes'],
                    'prestamos.plazo_credito': serializer.validated_data['plazo_credito'],
                    'variable_fecha_primer_pago': serializer.validated_data['variable_fecha_primer_pago'],
                    'variable_fecha_ultimo_pago': serializer.validated_data['variable_fecha_ultimo_pago'],
                }
                # Llamar a docx_replace con el documento y las sustituciones
                docx_replace(doc, **replacements)

                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="pagare_{serializer.validated_data["nombre_completo"]}.docx"'

                return response
            except Exception as e:
                return Response({'error': 'Error al procesar el documento.', 'details': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    


    
@api_view(['PATCH'])
@parser_classes([MultiPartParser, FormParser])
def update_cliente(request, cliente_id):
    cliente = get_object_or_404(Cliente, pk=cliente_id)
    serializer = ClienteSerializer(cliente, data=request.data, partial=True)

    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=status.HTTP_200_OK)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


def download_image(request, cliente_id):
    cliente = get_object_or_404(Cliente, pk=cliente_id)
    if cliente.foto_identificacion:
        file_url = cliente.foto_identificacion.url
        try:
            # Usar requests para obtener el archivo desde la URL
            response = requests.get(file_url, stream=True)
            if response.status_code == 200:
                # Preparar una respuesta de archivo con el contenido de la URL
                return FileResponse(response.raw, as_attachment=True, filename=os.path.basename(file_url))
            else:
                raise Http404("File not found on server.")
        except Exception as e:
            raise Http404(f"File not found: {e}")
    else:
        raise Http404("No image found for this client.")


@api_view(['POST'])
def register(request):
    serializer = UserSerializer(data=request.data)

    if serializer.is_valid():
        serializer.save()   
        user = User.objects.get(username=serializer.data['username'])    
        user.set_password(serializer.data['password'])
        user.save()

        token = Token.objects.create(user=user)

        return Response({'token': token.key, 'user': serializer.data}, status=status.HTTP_201_CREATED)
    
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@api_view(['POST'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def profile(request):

    serializer = UserSerializer(instance=request.user)

    return Response(serializer.data, status=status.HTTP_200_OK)


@api_view(['POST'])
def login(request):

    user = get_object_or_404(User, username=request.data['username'])

    if not user.check_password(request.data['password']):
        return Response({'error': 'Credenciales incorrectas'}, status=status.HTTP_400_BAD_REQUEST)

    token, created = Token.objects.get_or_create(user=user)

    serializer = UserSerializer(instance=user)   

    return Response({'token': token.key, 'user': serializer.data}, status=status.HTTP_200_OK)

# Create your views here.
class PrestamoViewSet(viewsets.ModelViewSet):
    queryset = Prestamo.objects.all()
    serializer_class = PrestamoSerializer

class ClienteViewSet(viewsets.ModelViewSet):
    queryset = Cliente.objects.all()
    serializer_class = ClienteSerializer

@api_view(['POST'])
def get_user_by_id(request):
    user_id = request.data.get('id')
    if not user_id:
        return Response({'error': 'El campo "id" es requerido.'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        user = User.objects.get(id=user_id)
    except User.DoesNotExist:
        return Response({'error': 'Usuario no encontrado.'}, status=status.HTTP_404_NOT_FOUND)

    serializer = UserSerializer(user)
    return Response(serializer.data, status=status.HTTP_200_OK)

# Vista para listar los pagos de un préstamo específico
@api_view(['GET'])
def listar_pagos_prestamo(request, prestamo_id):
    prestamo = get_object_or_404(Prestamo, id=prestamo_id)
    pagos = Pago.objects.filter(prestamo=prestamo).order_by('fecha_pago_programada')
    serializer = PagoSerializer(pagos, many=True)
    return Response(serializer.data, status=status.HTTP_200_OK)

# Vista para registrar el pago de una cuota
@api_view(['PATCH'])
def registrar_pago(request, pago_id):
    pago = get_object_or_404(Pago, id=pago_id)
    if not pago.pagado:
        pago.pagado = True
        pago.fecha_pago_realizado = datetime.now().date()  # Registrar la fecha actual como la fecha de pago realizado
        pago.save()



        return Response({'message': 'Pago registrado exitosamente'}, status=status.HTTP_200_OK)
    return Response({'error': 'Este pago ya ha sido registrado como pagado'}, status=status.HTTP_400_BAD_REQUEST)

# Vista para obtener los detalles de un pago específico
@api_view(['GET'])
def detalles_pago(request, pago_id):
    pago = get_object_or_404(Pago, id=pago_id)
    serializer = PagoSerializer(pago)
    return Response(serializer.data, status=status.HTTP_200_OK)


# Vista para desregistrar un pago
@api_view(['PATCH'])
def desregistrar_pago(request, pago_id):
    pago = get_object_or_404(Pago, id=pago_id)
    if pago.pagado:
        pago.pagado = False
        pago.fecha_pago_realizado = None  # Limpiar la fecha de pago realizado
        pago.save()


        return Response({'message': 'Pago desregistrado exitosamente'}, status=status.HTTP_200_OK)
    return Response({'error': 'Este pago ya está desregistrado'}, status=status.HTTP_400_BAD_REQUEST)



@api_view(['PATCH'])
def actualizar_estatus_prestamo(request, prestamo_id):
    try:
        # Obtener el préstamo correspondiente
        prestamo = get_object_or_404(Prestamo, pk=prestamo_id)
        pagos_pendientes = Pago.objects.filter(prestamo=prestamo, pagado=False)
        hoy = timezone.now().date()

        # Determinar el estatus del préstamo
        if not pagos_pendientes.exists():
            estatus = 'A Tiempo'
        elif any(pago.fecha_pago_programada < hoy for pago in pagos_pendientes):
            estatus = 'Atrasado'
        else:
            estatus = 'A Tiempo'

        # Actualizar el estatus del préstamo
        prestamo.estatus_pago = estatus
        prestamo.save()

        return Response({'estatus': estatus}, status=status.HTTP_200_OK)
    except Prestamo.DoesNotExist:
        return Response({'error': 'Préstamo no encontrado'}, status=status.HTTP_404_NOT_FOUND)



@api_view(['POST'])
def enviar_sms(request):
    # Credenciales de Twilio desde variables de entorno
    account_sid = config('TWILIO_ACCOUNT_SID')  # SID desde entorno
    auth_token = config('TWILIO_AUTH_TOKEN')    # Auth Token desde entorno
    client = Client(account_sid, auth_token)

    # Extraer parámetros de la solicitud
    mensaje = request.data.get('mensaje')
    cliente_id = request.data.get('cliente_id')

    # Validar que los datos estén presentes
    if not mensaje or not cliente_id:
        return Response(
            {"error": "Los campos 'mensaje' y 'cliente_id' son requeridos."},
            status=status.HTTP_400_BAD_REQUEST
        )

    # Obtener el cliente desde el modelo
    cliente = get_object_or_404(Cliente, id=cliente_id)

    # Validar que el cliente tenga un número de teléfono
    if not cliente.numero_telefono:
        return Response(
            {"error": "El cliente no tiene un número de teléfono registrado."},
            status=status.HTTP_400_BAD_REQUEST
        )

    # Formatear el número de teléfono con +52 si no lo tiene
    numero_telefono = cliente.numero_telefono.strip()
    if not numero_telefono.startswith('+52'):
        numero_telefono = f'+52{numero_telefono}'

    try:
        # Enviar el mensaje usando Twilio
        message = client.messages.create(
            body=mensaje,
            from_=config('TWILIO_PHONE_NUMBER'),  # Número de Twilio desde entorno
            to=numero_telefono  # Ahora el número está asegurado con el prefijo +52
        )

        # Retornar la respuesta exitosa
        return Response(
            {"message": "SMS enviado exitosamente.", "sms_sid": message.sid},
            status=status.HTTP_200_OK
        )
    except Exception as e:
        # Manejar errores durante el envío del SMS
        return Response(
            {"error": "Error al enviar el SMS.", "details": str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['PATCH'])
def actualizar_estado_prestamo(request, prestamo_id):
    """
    Cambia el estado de un préstamo entre ACTIVO y FINALIZADO.
    - Si el estado cambia a FINALIZADO, todos los pagos se marcan como pagados.
    - Si el estado cambia a ACTIVO, todos los pagos se marcan como no pagados.
    """
    prestamo = get_object_or_404(Prestamo, id=prestamo_id)
    nuevo_estado = request.data.get('estado')

    if nuevo_estado not in ['ACTIVO', 'FINALIZADO']:
        return Response({'error': 'Estado no válido. Debe ser ACTIVO o FINALIZADO.'}, status=status.HTTP_400_BAD_REQUEST)

    if prestamo.estado == nuevo_estado:
        return Response({'message': f'El préstamo ya está en estado {nuevo_estado}.'}, status=status.HTTP_400_BAD_REQUEST)

    # Actualizar el estado del préstamo
    prestamo.estado = nuevo_estado
    prestamo.save()

    # Actualizar los pagos asociados según el nuevo estado
    pagos_asociados = Pago.objects.filter(prestamo=prestamo)

    if nuevo_estado == 'FINALIZADO':
        pagos_asociados.update(pagado=True)
        return Response({
            'message': 'El préstamo ha sido finalizado y todos los pagos se marcaron como pagados.',
            'prestamo_estado': prestamo.estado,
            'pagos_actualizados': pagos_asociados.count()
        }, status=status.HTTP_200_OK)

    elif nuevo_estado == 'ACTIVO':
        pagos_asociados.update(pagado=False)
        return Response({
            'message': 'El préstamo ha sido reactivado y todos los pagos se marcaron como no pagados.',
            'prestamo_estado': prestamo.estado,
            'pagos_actualizados': pagos_asociados.count()
        }, status=status.HTTP_200_OK)
